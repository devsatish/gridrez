import io
from PyPDF2 import PdfReader
from docx import Document
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import PydanticOutputParser
from langchain_core.exceptions import OutputParserException
from app.models.schemas import ParsedResumeData, ResumeSummary, Education, SocialHandles


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text content from PDF bytes"""
    try:
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page in reader.pages:
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            except Exception as e:
                # Continue with other pages if one fails
                print(f"Warning: Failed to extract text from a PDF page: {str(e)}")
                continue

        if not text_parts:
            raise ValueError("No text could be extracted from the PDF file")
        
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text content from TXT bytes"""
    try:
        return file_content.decode("utf-8")
    except UnicodeDecodeError:
        # Try with error handling for non-UTF-8 files
        try:
            return file_content.decode("utf-8", errors="replace")
        except Exception as e:
            raise ValueError(f"Failed to decode text file: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text content from DOCX bytes"""
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text_parts = []
        try:
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
        except Exception as e:
            print(f"Warning: Failed to extract some paragraphs: {str(e)}")
        
        # Also extract text from tables
        try:
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
        except Exception as e:
            print(f"Warning: Failed to extract some table data: {str(e)}")
        
        if not text_parts:
            raise ValueError("No text could be extracted from the DOCX file")
        
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text(file_content: bytes, file_name: str) -> str:
    """Extract text based on file type"""
    file_name_lower = file_name.lower()
    if file_name_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_content)
    elif file_name_lower.endswith(".txt"):
        return extract_text_from_txt(file_content)
    elif file_name_lower.endswith(".docx"):
        return extract_text_from_docx(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_name}. Supported types: PDF, TXT, DOCX")


async def parse_resume(resume_text: str, resume_id: str) -> ResumeSummary:
    """Parse resume text using LangChain and GPT-4"""

    parser = PydanticOutputParser(pydantic_object=ParsedResumeData)

    prompt = ChatPromptTemplate.from_messages([
        ("system", """You are an expert resume parser. Analyze the provided text and extract structured information from a resume/CV.

IMPORTANT RULES:
1. Only parse if the text is clearly a resume or CV. A resume should contain:
   - A person's name
   - Work experience, job history, skills, education, or professional qualifications
2. If the text is NOT a resume (e.g., job description, article, general document), set name to null.
3. Handle partial or incomplete resumes gracefully:
   - If name is missing but it's clearly a resume, try to infer from context or use "Name not provided"
   - If currentRole is missing, use "Not specified"
   - If experienceYears cannot be determined, use 0
   - If skills are missing, return an empty list
   - If education is missing, return an empty list
   - If summary cannot be generated, use "No summary available"
   - If email/phone/location/social handles are missing, set to null (not empty strings)
4. Extract contact information carefully:
   - Email: Extract full email addresses (e.g., john.doe@example.com)
   - Phone: Extract phone numbers in any format (e.g., +1-555-123-4567, (555) 123-4567, etc.)
   - Location: Extract city, state/province, country (e.g., "San Francisco, CA, USA" or "London, UK")
   - Social Handles: Extract URLs or usernames for LinkedIn, GitHub, Twitter/X, portfolio websites
     * LinkedIn: Look for linkedin.com/in/ URLs or LinkedIn mentions
     * GitHub: Look for github.com/ URLs or GitHub usernames
     * Twitter/X: Look for twitter.com/ or x.com/ URLs or @handles
     * Portfolio: Look for personal websites or portfolio URLs
5. Be accurate and thorough. If information is not explicitly stated, make reasonable inferences based on context.
6. Always return valid data - never leave required fields as null unless the document is clearly not a resume.

{format_instructions}"""),
        ("user", """Please parse the following text and extract the requested information. 

If this is NOT a resume or CV, set the name field to null.
If this IS a resume but some information is missing, use the default values specified in the instructions.

---
{resume_text}
---

Extract all relevant information following the specified format. Handle missing data gracefully.""")
    ])

    llm = ChatOpenAI(model="gpt-4", temperature=0)

    chain = prompt | llm | parser

    try:
        parsed_data: ParsedResumeData = await chain.ainvoke({
            "resume_text": resume_text,
            "format_instructions": parser.get_format_instructions()
        })
    except OutputParserException as e:
        # Check if the error is due to null name (indicating non-resume)
        error_str = str(e)
        if "name" in error_str.lower() and ("null" in error_str.lower() or "none" in error_str.lower()):
            raise ValueError(
                "The uploaded file does not appear to be a resume or CV. "
                "Please upload a document containing a person's professional background, work experience, skills, and education."
            )
        # Re-raise with a more user-friendly message
        raise ValueError(
            "Unable to parse the resume. The file may not be a valid resume format. "
            "Please ensure the file contains clear resume information including name, work experience, and skills."
        )

    # Validate that we got a valid name (required to confirm it's a resume)
    if not parsed_data.name or parsed_data.name.strip() == "" or parsed_data.name.lower() == "null":
        raise ValueError(
            "The uploaded file does not appear to be a resume or CV. "
            "A resume should contain a person's name, work experience, skills, and education. "
            "Please upload a valid resume document."
        )

    # Handle partial data gracefully - provide defaults for missing fields
    name = parsed_data.name.strip() if parsed_data.name else "Name not provided"
    current_role = parsed_data.currentRole.strip() if parsed_data.currentRole else "Not specified"
    experience_years = parsed_data.experienceYears if parsed_data.experienceYears is not None else 0
    skills = parsed_data.skills if parsed_data.skills else []
    education = parsed_data.education if parsed_data.education else []
    summary_text = parsed_data.summary.strip() if parsed_data.summary else "No summary available"
    
    # Handle contact information - normalize but keep None if not found
    email = parsed_data.email.strip() if parsed_data.email and parsed_data.email.strip() else None
    phone = parsed_data.phone.strip() if parsed_data.phone and parsed_data.phone.strip() else None
    location = parsed_data.location.strip() if parsed_data.location and parsed_data.location.strip() else None
    
    # Handle social handles - create SocialHandles object if any handles exist
    social_handles = None
    try:
        if parsed_data.socialHandles:
            sh = parsed_data.socialHandles
            # Only create SocialHandles if at least one handle exists
            if (sh.linkedin or sh.github or sh.twitter or sh.portfolio or 
                (sh.other and len(sh.other) > 0)):
                try:
                    social_handles = SocialHandles(
                        linkedin=sh.linkedin.strip() if sh.linkedin and sh.linkedin.strip() else None,
                        github=sh.github.strip() if sh.github and sh.github.strip() else None,
                        twitter=sh.twitter.strip() if sh.twitter and sh.twitter.strip() else None,
                        portfolio=sh.portfolio.strip() if sh.portfolio and sh.portfolio.strip() else None,
                        other=[s.strip() for s in sh.other if s and s.strip()] if sh.other else []
                    )
                except Exception as e:
                    # If SocialHandles creation fails, just skip it
                    print(f"Warning: Failed to create SocialHandles: {str(e)}")
                    social_handles = None
    except Exception as e:
        # If accessing socialHandles fails, just skip it
        print(f"Warning: Failed to process social handles: {str(e)}")
        social_handles = None

    # Convert to ResumeSummary with normalized data
    try:
        summary = ResumeSummary(
            id=resume_id,
            name=name,
            currentRole=current_role,
            experienceYears=experience_years,
            skills=skills,
            education=education,
            summary=summary_text,
            email=email,
            phone=phone,
            location=location,
            socialHandles=social_handles
        )
    except Exception as e:
        # Fallback: try creating without optional fields
        try:
            summary = ResumeSummary(
                id=resume_id,
                name=name,
                currentRole=current_role,
                experienceYears=experience_years,
                skills=skills,
                education=education,
                summary=summary_text,
                email=None,
                phone=None,
                location=None,
                socialHandles=None
            )
        except Exception as fallback_error:
            # If even fallback fails, raise a clear error
            raise ValueError(
                f"Failed to create resume summary: {str(fallback_error)}. "
                "Please ensure the resume contains valid information."
            )

    return summary
